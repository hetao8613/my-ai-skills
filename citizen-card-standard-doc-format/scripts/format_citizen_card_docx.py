#!/usr/bin/env python3
"""Normalize a DOCX to the Citizen Card standard meeting-document format."""

from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


CN_NUMS = "一二三四五六七八九十十一十二十三十四十五十六十七十八十九二十"
H1_RE = re.compile(r"^[一二三四五六七八九十]{1,3}、")
H2_RE = re.compile(r"^（[一二三四五六七八九十]{1,3}）")
H3_RE = re.compile(r"^\d+[.．、]")
H4_RE = re.compile(r"^（\d+）")


def set_east_asian_font(run, font_name: str, size_pt: float | None = None, bold: bool | None = None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, line_pt: float = 29, first_line: bool = False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(line_pt)
    fmt.first_line_indent = Pt(32) if first_line else Pt(0)


def remove_numbering(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def replace_prefix(text: str, prefix: str, pattern: re.Pattern[str]) -> str:
    stripped = text.strip()
    if pattern.match(stripped):
        return pattern.sub(prefix, stripped, count=1)
    return prefix + stripped


def classify_paragraph(paragraph, index: int) -> str:
    text = paragraph.text.strip()
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    if not text:
        return "blank"
    if index == 0 or "title" in style_name and len(text) < 80:
        return "title"
    if "heading 1" in style_name or H1_RE.match(text):
        return "h1"
    if "heading 2" in style_name or H2_RE.match(text):
        return "h2"
    if "heading 3" in style_name or H3_RE.match(text):
        return "h3"
    if "heading 4" in style_name or H4_RE.match(text):
        return "h4"
    return "body"


def rewrite_paragraph_text(paragraph, text: str):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def style_paragraphs(doc: Document, renumber: bool = True):
    h1_count = 0
    h2_count = 0
    body_index = 0

    for paragraph in doc.paragraphs:
        kind = classify_paragraph(paragraph, body_index)
        if kind != "blank":
            body_index += 1

        if kind == "blank":
            set_paragraph_format(paragraph)
            continue

        remove_numbering(paragraph)

        if kind == "title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(paragraph)
            for run in paragraph.runs:
                set_east_asian_font(run, "方正小标宋简体", 22, False)
            continue

        if kind == "h1":
            h1_count += 1
            h2_count = 0
            if renumber and h1_count <= len(CN_NUMS):
                rewrite_paragraph_text(paragraph, replace_prefix(paragraph.text, f"{CN_NUMS[h1_count - 1]}、", H1_RE))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(paragraph)
            for run in paragraph.runs:
                set_east_asian_font(run, "黑体", 16, False)
            continue

        if kind == "h2":
            h2_count += 1
            if renumber and h2_count <= len(CN_NUMS):
                rewrite_paragraph_text(paragraph, replace_prefix(paragraph.text, f"（{CN_NUMS[h2_count - 1]}）", H2_RE))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(paragraph)
            for run in paragraph.runs:
                set_east_asian_font(run, "楷体_GB2312", 16, False)
            continue

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_format(paragraph, first_line=kind == "body")
        font = "仿宋_GB2312"
        size = 16
        for run in paragraph.runs:
            set_east_asian_font(run, font, size, False)


def set_cell_margin(cell, margin_dxa: int = 60):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_dxa))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def allow_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is not None:
        tr_pr.remove(cant_split)


def style_tables(doc: Document, strict_table_font: bool = False):
    table_size = 16 if strict_table_font else 12
    line_pt = 29 if strict_table_font else 20
    for table in doc.tables:
        if table.rows:
            repeat_table_header(table.rows[0])
        for row_idx, row in enumerate(table.rows):
            allow_row_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margin(cell)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    set_paragraph_format(paragraph, line_pt=line_pt, first_line=False)
                    for run in paragraph.runs:
                        set_east_asian_font(run, "仿宋_GB2312", table_size, row_idx == 0)


def set_doc_grid(section):
    sect_pr = section._sectPr
    doc_grid = sect_pr.find(qn("w:docGrid"))
    if doc_grid is None:
        doc_grid = OxmlElement("w:docGrid")
        sect_pr.append(doc_grid)
    doc_grid.set(qn("w:type"), "linesAndChars")
    doc_grid.set(qn("w:linePitch"), "580")


def setup_sections(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.6)
        section.bottom_margin = Cm(2.8)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(2.6)
        section.footer_distance = Cm(2.6)
        section.different_first_page_header_footer = False
        section.odd_and_even_pages_header_footer = True
        set_doc_grid(section)


def page_field_run(paragraph):
    run = paragraph.add_run("-")
    set_east_asian_font(run, "宋体", 14, False)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    tail = paragraph.add_run("-")
    set_east_asian_font(tail, "宋体", 14, False)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def setup_footers(doc: Document):
    settings = doc.settings.element
    even = settings.find(qn("w:evenAndOddHeaders"))
    if even is None:
        even = OxmlElement("w:evenAndOddHeaders")
        settings.append(even)

    for section in doc.sections:
        section.odd_and_even_pages_header_footer = True
        footers = (
            (section.footer, WD_ALIGN_PARAGRAPH.RIGHT),
            (section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT),
        )
        for footer, align in footers:
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            clear_paragraph(paragraph)
            paragraph.alignment = align
            set_paragraph_format(paragraph, line_pt=14, first_line=False)
            page_field_run(paragraph)


def normalize_docx(input_path: Path, output_path: Path, strict_table_font: bool = False, renumber: bool = True):
    doc = Document(str(input_path))
    setup_sections(doc)
    style_paragraphs(doc, renumber=renumber)
    style_tables(doc, strict_table_font=strict_table_font)
    setup_footers(doc)
    doc.save(str(output_path))


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Source .docx file")
    parser.add_argument("--output", "-o", type=Path, help="Output .docx file")
    parser.add_argument("--strict-table-font", action="store_true", help="Keep table text at 16 pt")
    parser.add_argument("--no-manual-renumber", action="store_true", help="Do not rewrite heading numbering")
    args = parser.parse_args()

    if args.input.suffix.lower() != ".docx":
        raise SystemExit("Input must be a .docx file")
    output = args.output or args.input.with_name(f"{args.input.stem}-格式调整版.docx")
    normalize_docx(args.input, output, args.strict_table_font, not args.no_manual_renumber)
    print(output)


if __name__ == "__main__":
    main()
